from docx import Document
from pathlib import Path
import zipfile
import uuid
import os

# Автоматическое определение пути к папке проекта
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"
GENERATED_DIR = Path("generated")
GENERATED_DIR.mkdir(exist_ok=True)

def replace_text(doc, data: dict):
    """Рекурсивная замена текста в параграфах и таблицах"""
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                for run in p.runs:
                    # Заменяем текст, обрабатывая None как пустую строку
                    run.text = run.text.replace(key, str(value if value is not None else ""))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Рекурсивно заходим в параграфы каждой ячейки
                replace_text(cell, data)

def build_template_data(student_row: dict) -> dict:
    """
    Сопоставляет чистые колонки из БД (student_row) с метками {{...}} в Word.
    """
    return {
        "{{ФИО_обучающегося}}": student_row.get("ФИО_обучающегося"),
        "{{Дата_рождения}}": student_row.get("Дата_рождения"),
        "{{наименование_модуля}}": student_row.get("наименование_модуля"),
        "{{день_начала_ПП}}": student_row.get("день_начала_ПП"),
        "{{месяц_начала_ПП}}": student_row.get("месяц_начала_ПП"),
        "{{год_начала_ПП}}": student_row.get("год_начала_ПП"),
        "{{день_конца_ПП}}": student_row.get("день_конца_ПП"),
        "{{месяц_конца_ПП}}": student_row.get("месяц_конца_ПП"),
        "{{год_конца_ПП}}": student_row.get("год_конца_ПП"),
        "{{код_специальности}}": student_row.get("код_специальности"),
        "{{наименование_специальности}}": student_row.get("наименование_специальности"),
        "{{количество_часов_ПП}}": student_row.get("количество_часов_ПП"),
        "{{ФИО_преподавателя}}": student_row.get("ФИО_преподавателя"),
        "{{тел_преподаваттеля}}": student_row.get("тел_преподаваттеля"),
        "{{название_организации}}": student_row.get("название_организации"),
        "{{адрес_организации}}": student_row.get("адрес_организации"),
        "{{Наименование_помещений}}": student_row.get("Наименование_помещений"),
        "{{должность_отв_организации}}": student_row.get("должность_отв_организации"),
        "{{ФИО_отв_организации}}": student_row.get("ФИО_отв_организации"),
        "{{тел_отв_организации}}": student_row.get("тел_отв_организации"),
        "{{инициалы_отв_организации}}": student_row.get("инициалы_отв_организации"),
        "{{должность_ответственного}}": student_row.get("должность_ответственного"),
        "{{ФИО_ответственного}}": student_row.get("ФИО_ответственного"),
        "{{тел_ответственного}}": student_row.get("тел_ответственного"),
    }

def generate_all_docs(student_data: dict, selected_files: list = None):
    """
    Основная логика генерации ZIP-архива.
    Принимает словарь данных и список имен файлов для обработки.
    """
    all_templates = [
        "Отчёт производственная.docx",
        "Отчёт учебная.docx",
        "Аттестационный лист производственная.docx",
        "Аттестационный лист учебная.docx",
    ]

    # Если список не передан, берем все. Если передан — только то, что в списке.
    templates_to_process = selected_files if selected_files else all_templates

    folder_name = str(uuid.uuid4())[:8]
    work_dir = GENERATED_DIR / folder_name
    work_dir.mkdir(exist_ok=True)
    
    generated_files = []
    
    print(f"\n--- ГЕНЕРАЦИЯ В ПАПКЕ: {work_dir} ---")
    
    for tpl in templates_to_process:
        tpl_path = TEMPLATES_DIR / tpl
        out_file = work_dir / tpl
        
        if not tpl_path.exists():
            print(f"❌ ФАЙЛ ШАБЛОНА НЕ НАЙДЕН: {tpl_path}")
            continue
            
        try:
            doc = Document(tpl_path)
            replace_text(doc, student_data) 
            doc.save(out_file)
            generated_files.append(out_file)
            print(f"✅ УСПЕШНО СОЗДАН: {tpl}")
        except Exception as e:
            print(f"🔥 ОШИБКА ПРИ ОБРАБОТКЕ {tpl}: {e}")

    if not generated_files:
        return None

    zip_filename = f"docs_{folder_name}.zip"
    zip_path = GENERATED_DIR / zip_filename
    
    with zipfile.ZipFile(zip_path, "w") as z:
        for f in generated_files:
            z.write(f, arcname=f.name)
            
    return zip_path
